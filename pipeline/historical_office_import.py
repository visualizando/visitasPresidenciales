from __future__ import annotations

import hashlib
import re
from collections import defaultdict
from collections.abc import Iterable
from datetime import date, datetime, time, timedelta
from pathlib import Path
from typing import Any
from urllib.parse import quote

import openpyxl
import pymupdf
import xlrd
from docx import Document

from pipeline.build_web import build_web_data
from pipeline.models import AccessRecord
from pipeline.normalize import canonical_name, document_identity, entity_id, fold_text, stable_id
from pipeline.storage import load_json, utc_now, write_json_atomic, write_partition

MONTHS = {
    "ENERO": 1,
    "ENE": 1,
    "FEBRERO": 2,
    "FEB": 2,
    "MARZO": 3,
    "MAR": 3,
    "ABRIL": 4,
    "ABR": 4,
    "MAYO": 5,
    "MAY": 5,
    "JUNIO": 6,
    "JUN": 6,
    "JULIO": 7,
    "JUL": 7,
    "AGOSTO": 8,
    "AGO": 8,
    "SEPTIEMBRE": 9,
    "SETIEMBRE": 9,
    "SEP": 9,
    "OCTUBRE": 10,
    "OCT": 10,
    "NOVIEMBRE": 11,
    "NOV": 11,
    "DICIEMBRE": 12,
    "DIC": 12,
}
OLIVOS_YEARS = {2016, 2017, 2018, 2019}


def import_historical_office(
    source_root: Path,
    data_dir: Path,
    web_data_dir: Path,
    *,
    olivos_years: set[int] | None = None,
    include_casa_2020: bool = True,
    rebuild_web: bool = True,
) -> dict[str, int]:
    """Import structured historical XLS/XLSX/DOCX sources without retaining originals in Git."""
    years = olivos_years or OLIVOS_YEARS
    source_root = source_root.resolve()
    candidates = _candidate_files(source_root, years, include_casa_2020)
    manifest = load_json(
        data_dir / "manifest.json", {"version": 1, "generated_at": None, "files": {}}
    )
    partitions_root = data_dir / "partitions"
    diagnostics: defaultdict[str, int] = defaultdict(int)
    imported = 0
    changed = False

    for source in candidates:
        relative = source.relative_to(source_root).as_posix()
        source_path = f"raw/{relative}"
        source_id = stable_id("src_", source_path)
        info = source.stat()
        previous = manifest["files"].get(source_path)
        if _entry_matches_metadata(previous, info, partitions_root):
            diagnostics["unchanged_sources"] += 1
            continue

        checksum = _sha256(source)
        if _entry_is_current(previous, checksum, partitions_root):
            previous["size"] = info.st_size
            previous["last_modified"] = str(info.st_mtime_ns)
            diagnostics["unchanged_by_checksum"] += 1
            changed = True
            continue

        changed = True
        for old_partition in (
            partitions_root.rglob(f"{source_id}.parquet") if partitions_root.exists() else []
        ):
            old_partition.unlink()

        location = "casa-rosada" if "CASA ROSADA" in fold_text(relative) else "olivos"
        suffix = source.suffix.lower()
        status = "active"
        quarantine_reason: str | None = None
        try:
            if suffix == ".xlsx":
                records = _parse_xlsx(source, source_path, source_id, location, diagnostics)
                parser = (
                    "casa-rosada-xlsx-historico-v2"
                    if location == "casa-rosada"
                    else "olivos-xlsx-historico-v2"
                )
            elif suffix == ".xls":
                records = _parse_xls(source, source_path, source_id, diagnostics)
                parser = "olivos-xls-historico-v1"
            elif suffix == ".docx":
                records = _parse_docx(source, source_path, source_id, diagnostics)
                parser = "olivos-docx-historico-v2"
            else:
                records, parser, quarantine_reason = _inspect_pdf(source)
                status = "quarantined"
        except Exception as error:
            records = []
            parser = f"olivos-{suffix.lstrip('.') or 'archivo'}-error-v1"
            status = "quarantined"
            quarantine_reason = "archivo_danado"
            diagnostics[f"errors_{type(error).__name__}"] += 1

        if status == "active" and not records:
            empty_reason = _empty_source_reason(source)
            if empty_reason == "sin_novedad":
                diagnostics["sources_without_activity"] += 1
            else:
                status = "quarantined"
                quarantine_reason = empty_reason
                diagnostics[f"quarantined_{empty_reason}"] += 1

        grouped: dict[tuple[int, int], list[AccessRecord]] = defaultdict(list)
        for record in records:
            timestamp = record.occurred_at or record.entered_at or record.exited_at
            if timestamp:
                grouped[(timestamp.year, timestamp.month)].append(record)
        partition_paths: list[str] = []
        for (year, month), rows in sorted(grouped.items()):
            partition = Path(location) / str(year) / f"{month:02d}" / f"{source_id}.parquet"
            write_partition(partitions_root / partition, rows)
            partition_paths.append(partition.as_posix())

        year, month = _source_period(source)
        manifest["files"][source_path] = {
            "source_id": source_id,
            "url": "local-source:///" + quote(source_path),
            "path": source_path,
            "location": location,
            "year": year,
            "month": month,
            "size": info.st_size,
            "etag": None,
            "last_modified": str(info.st_mtime_ns),
            "sha256": checksum,
            "parser": parser,
            "record_count": len(records),
            "status": status,
            "partition": None,
            "partitions": partition_paths,
            "processed_at": utc_now(),
        }
        if quarantine_reason:
            manifest["files"][source_path]["quarantine_reason"] = quarantine_reason
        diagnostics["sources"] += 1
        if not records:
            diagnostics["sources_without_records"] += 1
        imported += len(records)

    if changed:
        manifest["generated_at"] = utc_now()
        write_json_atomic(data_dir / "manifest.json", manifest)
    web = build_web_data(data_dir, web_data_dir) if rebuild_web and changed else {}
    return {"imported": imported, **dict(sorted(diagnostics.items())), **web}


def _entry_is_current(entry: Any, checksum: str, partitions_root: Path) -> bool:
    return _entry_has_current_output(entry, partitions_root) and entry.get("sha256") == checksum


def _entry_matches_metadata(entry: Any, info: Any, partitions_root: Path) -> bool:
    return (
        _entry_has_current_output(entry, partitions_root)
        and entry.get("size") == info.st_size
        and entry.get("last_modified") == str(info.st_mtime_ns)
    )


def _entry_has_current_output(entry: Any, partitions_root: Path) -> bool:
    if not isinstance(entry, dict):
        return False
    if entry.get("status") not in {"active", "quarantined"}:
        return False

    parser = str(entry.get("parser") or "")
    current_parsers = {
        "casa-rosada-xlsx-historico-v2",
        "olivos-xlsx-historico-v2",
        "olivos-xls-historico-v1",
        "olivos-docx-historico-v2",
        "pdf-sin-texto-extraible-v1",
        "pdf-historico-requiere-parser-v1",
        "olivos-xlsx-error-v1",
        "olivos-xls-error-v1",
        "olivos-docx-error-v1",
        "olivos-pdf-error-v1",
    }
    if parser not in current_parsers:
        return False

    partitions = entry.get("partitions") or []
    if not partitions and entry.get("partition"):
        partitions = [entry["partition"]]
    return all((partitions_root / relative).is_file() for relative in partitions)


def _candidate_files(root: Path, years: set[int], include_casa_2020: bool) -> list[Path]:
    candidates: list[Path] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.name.startswith("~$") or path.stat().st_size < 1000:
            continue
        suffix = path.suffix.lower()
        relative = fold_text(path.relative_to(root).as_posix())
        path_year, _ = _source_period(path)
        if (
            "QUINTA DE OLIVOS" in relative
            and path_year in years
            and suffix in {".xlsx", ".xls", ".docx", ".pdf"}
        ) or (
            include_casa_2020
            and "CASA ROSADA" in relative
            and path_year == 2020
            and suffix == ".xlsx"
        ):
            candidates.append(path)
    return candidates


def _parse_xls(
    source: Path,
    source_path: str,
    source_id: str,
    diagnostics: defaultdict[str, int],
) -> list[AccessRecord]:
    workbook = xlrd.open_workbook(source, on_demand=True)
    records: list[AccessRecord] = []
    try:
        for sheet_number, worksheet in enumerate(workbook.sheets(), 1):
            rows = [
                [_xls_cell_value(worksheet.cell(row, column), workbook.datemode) for column in range(worksheet.ncols)]
                for row in range(worksheet.nrows)
            ]
            records.extend(
                _parse_olivos_rows(
                    rows,
                    worksheet.name,
                    source,
                    source_path,
                    source_id,
                    sheet_number,
                    diagnostics,
                )
            )
    finally:
        workbook.release_resources()
    return records


def _xls_cell_value(cell: Any, datemode: int) -> Any:
    if cell.ctype == xlrd.XL_CELL_DATE:
        return xlrd.xldate_as_datetime(cell.value, datemode)
    return cell.value


def _inspect_pdf(source: Path) -> tuple[list[AccessRecord], str, str]:
    document = pymupdf.open(source)
    try:
        text_characters = sum(len(page.get_text("text").strip()) for page in document)
    finally:
        document.close()
    if text_characters == 0:
        return [], "pdf-sin-texto-extraible-v1", "pdf_escaneado"
    return [], "pdf-historico-requiere-parser-v1", "formato_no_compatible"


def _parse_xlsx(
    source: Path,
    source_path: str,
    source_id: str,
    location: str,
    diagnostics: defaultdict[str, int],
) -> list[AccessRecord]:
    workbook = openpyxl.load_workbook(source, read_only=True, data_only=True)
    records: list[AccessRecord] = []
    try:
        for sheet_number, worksheet in enumerate(workbook.worksheets, 1):
            if location == "casa-rosada":
                records.extend(
                    _parse_casa_sheet(
                        worksheet, source, source_path, source_id, sheet_number, diagnostics
                    )
                )
            else:
                records.extend(
                    _parse_olivos_rows(
                        worksheet.iter_rows(values_only=True),
                        worksheet.title,
                        source,
                        source_path,
                        source_id,
                        sheet_number,
                        diagnostics,
                    )
                )
    finally:
        workbook.close()
    return records


def _parse_docx(
    source: Path,
    source_path: str,
    source_id: str,
    diagnostics: defaultdict[str, int],
) -> list[AccessRecord]:
    document = Document(source)
    records: list[AccessRecord] = []
    for table_number, table in enumerate(document.tables, 1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        records.extend(
            _parse_olivos_rows(
                rows,
                f"tabla-{table_number}",
                source,
                source_path,
                source_id,
                table_number,
                diagnostics,
            )
        )
    return records


def _parse_olivos_rows(
    rows: Iterable[Iterable[Any]],
    sheet_title: str,
    source: Path,
    source_path: str,
    source_id: str,
    source_page: int,
    diagnostics: defaultdict[str, int],
) -> list[AccessRecord]:
    materialized = [list(row) for row in rows]
    while materialized and not any(_clean_text(value) for value in materialized[-1]):
        materialized.pop()
    if not materialized:
        return []
    header_index, headers = _find_olivos_header(materialized)
    if header_index is None:
        return _parse_olivos_daily_rows(
            materialized,
            sheet_title,
            source,
            source_path,
            source_id,
            source_page,
            diagnostics,
        )
    name_index = _column(
        headers,
        "APELLIDO Y NOMBRE",
        "APELLIDOS Y NOMBRES",
        "NOMBRE Y APELLIDO",
        "AUTORIDAD/CONDUCTOR",
        "AUTORIDAD",
    )
    document_index = _column(headers, "DOCUMENTO", "DNI")
    destination_index = _column(headers, "CONCURRE A", "CONCURRE PARA", "DESTINO")
    authorized_index = _column(headers, "AUTORIZADO")
    entry_index = _column(headers, "HORA ENTRADA", "HORA DE ENTRADA", "INGRESO")
    exit_index = _column(headers, "HORA SALIDA", "HORA DE SALIDA", "EGRESO")
    if name_index is None or entry_index is None:
        return []

    default_date = _infer_date(source, sheet_title, materialized[:15])
    record_type = (
        "vehicle"
        if "VEH" in fold_text(sheet_title)
        or ("MARCA/MODELO" in headers and "AUDIENCIA" not in fold_text(sheet_title))
        else "person"
    )
    records: list[AccessRecord] = []
    for row in materialized[header_index + 1 :]:
        name = _cell(row, name_index)
        folded_name = fold_text(name)
        if not name or _is_non_person_name(name) or folded_name in {
            "APELLIDO Y NOMBRE",
            "APELLIDOS Y NOMBRES",
            "AUTORIDAD",
            "SIN NOVEDAD",
        }:
            continue
        if folded_name.startswith(("FUNCIONARIO", "TRABAJO", "VISITA", "OTRO")):
            continue
        entered = _cell_datetime(_value(row, entry_index), default_date)
        exited = _cell_datetime(_value(row, exit_index), default_date)
        if entered and exited and exited < entered:
            if exited.date() == entered.date():
                exited += timedelta(days=1)
            else:
                diagnostics["exit_before_entry"] += 1
                exited = None
        if not entered and not exited:
            continue
        document = _cell(row, document_index)
        destination = _clean(_cell(row, destination_index))
        authorized = _clean(_cell(row, authorized_index))
        activity_values = []
        if destination_index is not None and authorized_index is not None:
            activity_values = [
                _cell(row, index)
                for index in range(destination_index + 1, authorized_index)
                if _cell(row, index)
            ]
        records.append(
            _record(
                name=name,
                document=document,
                location="olivos",
                record_type=record_type,
                source_id=source_id,
                source_path=source_path,
                source_page=source_page,
                entered=entered,
                exited=exited,
                destination=destination,
                activity=" · ".join(activity_values) or None,
                authorized_by=authorized,
                raw_text=" | ".join(_clean_text(value) for value in row if _clean_text(value)),
            )
        )
    return records


def _parse_olivos_daily_rows(
    rows: list[list[Any]],
    sheet_title: str,
    source: Path,
    source_path: str,
    source_id: str,
    source_page: int,
    diagnostics: defaultdict[str, int],
) -> list[AccessRecord]:
    family = _daily_family(source_path, sheet_title)
    if family is None:
        return []
    default_date = _infer_date(source, sheet_title, rows[:15])
    if default_date is None:
        diagnostics["daily_missing_date"] += 1
        return []

    layout = _daily_layout(rows, family)
    if layout is None:
        _increment(diagnostics, f"daily_{family}_unknown_layout")
        return []
    start, name_index, document_index, destination_index, device_index = layout
    records: list[AccessRecord] = []
    occurred = default_date.replace(hour=0, minute=0, second=0, microsecond=0)
    for row in rows[start:]:
        name = _cell(row, name_index)
        if not _looks_person_name(name):
            continue
        document = _cell(row, document_index)
        destination = _clean(_cell(row, destination_index))
        device = _clean(_cell(row, device_index))
        records.append(
            _record(
                name=name,
                document=document,
                location="olivos",
                record_type="vehicle" if family == "vehicle" else "person",
                source_id=source_id,
                source_path=source_path,
                source_page=source_page,
                entered=None,
                exited=None,
                occurred=occurred,
                destination=destination,
                activity="Ingreso diario sin horario informado",
                authorized_by=None,
                device=device,
                raw_text=" | ".join(_clean_text(value) for value in row if _clean_text(value)),
            )
        )
    if records:
        _increment(diagnostics, f"daily_{family}_sheets")
        _increment(diagnostics, f"daily_{family}_records", len(records))
    return records


def _daily_family(source_path: str, sheet_title: str) -> str | None:
    context = fold_text(f"{source_path} {sheet_title}")
    if any(label in context for label in ("INGRESO EN MOVIL", "INGRESO MOVIL", "VEHICUL")):
        return "vehicle"
    if any(label in context for label in ("INGRESO A PIE", "MOVIMIENTO DE PERSONAS", " A PIE")):
        return "person"
    return None


def _daily_layout(
    rows: list[list[Any]], family: str
) -> tuple[int, int, int | None, int | None, int | None] | None:
    for index, row in enumerate(rows[:35]):
        headers = [fold_text(_clean_text(value)) for value in row]
        name_index = _column(
            headers,
            "APELLIDO Y NOMBRE",
            "APELLIDOS Y NOMBRES",
            "NOMBRE Y APELLIDO",
            "AUTORIDAD/CONDUCTOR",
            "AUTORIDAD",
        )
        if name_index is None:
            continue
        return (
            index + 1,
            name_index,
            _column(headers, "DOCUMENTO", "DNI"),
            _column(headers, "CONCURRE A", "CONCURRE PARA", "DESTINO", "VISITA A"),
            _column(headers, "PATENTE", "DOMINIO", "MARCA/MODELO"),
        )

    for index, row in enumerate(rows[:35]):
        if family == "vehicle":
            if len(row) < 2 or not _looks_person_name(_cell(row, 1)):
                continue
            document_index = 2 if _looks_document_value(_value(row, 2)) else None
            destination_index = 3 if document_index is not None else 2
            return index, 1, document_index, destination_index, 0
        if not row or not _looks_person_name(_cell(row, 0)):
            continue
        document_index = 1 if _looks_document_value(_value(row, 1)) else None
        destination_index = 2 if document_index is not None else 1
        return index, 0, document_index, destination_index, None
    return None


def _looks_person_name(value: str) -> bool:
    folded = fold_text(value)
    if _is_non_person_name(value) or any(
        marker in folded
        for marker in (
            "APELLIDO",
            "AUTORIDAD",
            "CONDUCTOR",
            "PLANILLA",
            "RESIDENCIA",
            "DIRECCION DE SEGURIDAD",
            "SIN NOVEDAD",
            "PERSONAL AUTORIZADO",
        )
    ):
        return False
    tokens = folded.split()
    return len(tokens) >= 2 and sum(character.isalpha() for character in folded) >= 5


def _is_non_person_name(value: str) -> bool:
    folded = fold_text(value)
    if not folded or folded == "NAME" or folded[0].isdigit():
        return True
    if any(
        marker in folded
        for marker in (
            "REFORMA UNIVERSITARIA",
            "SIN NOVEDAD",
            "PERSONAL AUTORIZADO",
        )
    ):
        return True
    compact = folded.replace(" ", "")
    return bool(re.fullmatch(r"(?:[A-Z]{3}\d{3}|[A-Z]{2}\d{3}[A-Z]{2})", compact))


def _looks_document_value(value: Any) -> bool:
    digits = re.sub(r"\D", "", _clean_text(value))
    return 7 <= len(digits) <= 11


def _parse_casa_sheet(
    worksheet: Any,
    source: Path,
    source_path: str,
    source_id: str,
    source_page: int,
    diagnostics: defaultdict[str, int],
) -> list[AccessRecord]:
    rows = [list(row) for row in worksheet.iter_rows(values_only=True)]
    default_date = _infer_date(source, worksheet.title, rows[:20])
    if not default_date:
        return []
    header_index: int | None = None
    indices: tuple[int, int, int | None, int | None, int | None] | None = None
    for index, row in enumerate(rows[:25]):
        headers = [fold_text(_clean_text(value)) for value in row]
        name_index = _column(headers, "APELLIDO Y NOMBRE")
        entry_index = _column(headers, "ENTRO")
        if name_index is not None and entry_index is not None:
            header_index = index
            indices = (
                name_index,
                entry_index,
                _column(headers, "SALIO"),
                _column(headers, "ORGANISMO", "ORGANIZACION"),
                _column(headers, "DNI", "D N I"),
            )
            break
    if indices is None:
        indices = _infer_casa_indices(rows[:30])
        header_index = -1
    if indices is None:
        return []

    name_index, entry_index, exit_index, destination_index, document_index = indices
    records: list[AccessRecord] = []
    for row in rows[(header_index or 0) + 1 :]:
        name = _cell(row, name_index)
        if not name or fold_text(name) == "APELLIDO Y NOMBRE":
            continue
        entered = _cell_datetime(_value(row, entry_index), default_date)
        exited = _cell_datetime(_value(row, exit_index), default_date)
        if entered and exited and exited < entered:
            exited += timedelta(days=1)
        if not entered and not exited:
            continue
        records.append(
            _record(
                name=name,
                document=_cell(row, document_index),
                location="casa-rosada",
                record_type="visitor",
                source_id=source_id,
                source_path=source_path,
                source_page=source_page,
                entered=entered,
                exited=exited,
                destination=_clean(_cell(row, destination_index)),
                activity=None,
                authorized_by=None,
                raw_text=" | ".join(_clean_text(value) for value in row if _clean_text(value)),
            )
        )
    diagnostics["casa_sheets"] += 1
    return records


def _find_olivos_header(rows: list[list[Any]]) -> tuple[int | None, list[str]]:
    for index, row in enumerate(rows[:35]):
        headers = [fold_text(_clean_text(value)) for value in row]
        joined = " | ".join(headers)
        has_name = any(
            label in joined
            for label in (
                "APELLIDO Y NOMBRE",
                "APELLIDOS Y NOMBRES",
                "NOMBRE Y APELLIDO",
                "AUTORIDAD/CONDUCTOR",
            )
        ) or ("AUTORIDAD" in joined and ("HORA ENTRADA" in joined or "INGRESO" in joined))
        has_entry = any(label in joined for label in ("HORA ENTRADA", "HORA DE ENTRADA", "INGRESO"))
        if has_name and has_entry:
            return index, headers
    return None, []


def _infer_casa_indices(
    rows: list[list[Any]],
) -> tuple[int, int, int | None, int | None, int | None] | None:
    for row in rows:
        for index, value in enumerate(row):
            if index + 4 >= len(row) or not _looks_sequence(value):
                continue
            if _looks_clock(row[index + 1]) and canonical_name(_clean_text(row[index + 2])):
                return index + 2, index + 1, index + 3, index + 4, index + 5
    return None


def _record(
    *,
    name: str,
    document: str | None,
    location: str,
    record_type: str,
    source_id: str,
    source_path: str,
    source_page: int,
    entered: datetime | None,
    exited: datetime | None,
    destination: str | None,
    activity: str | None,
    authorized_by: str | None,
    raw_text: str,
    occurred: datetime | None = None,
    device: str | None = None,
) -> AccessRecord:
    clean_name = canonical_name(name)
    document_type, document_number, _ = document_identity(document)
    person_id = entity_id(clean_name, document_number)
    logical_time = occurred or entered or exited
    record_id = stable_id(
        "rec_",
        person_id,
        location,
        record_type,
        logical_time.isoformat() if logical_time else "",
        exited.isoformat() if exited else "",
        None,
        None,
        destination,
    )
    return AccessRecord(
        record_id=record_id,
        entity_id=person_id,
        canonical_name=clean_name,
        document_type=document_type,
        document_number=document_number,
        location=location,  # type: ignore[arg-type]
        record_type=record_type,  # type: ignore[arg-type]
        source_id=source_id,
        source_url="local-source:///" + quote(source_path),
        source_path=source_path,
        source_page=source_page,
        occurred_at=occurred,
        entered_at=entered,
        exited_at=exited,
        device=device,
        destination=destination,
        activity=activity,
        authorized_by=authorized_by,
        quality="high" if entered and exited else "medium",
        raw_text=raw_text,
    )


def _infer_date(source: Path, sheet_title: str, rows: list[list[Any]]) -> datetime | None:
    path_year, path_month = _source_period(source)
    texts = [sheet_title, source.stem]
    texts.extend(" ".join(_clean_text(value) for value in row if value is not None) for row in rows)
    for text in texts:
        numeric_text = _clean_text(text).upper()
        folded = fold_text(text)
        years = [int(value) for value in re.findall(r"20\d{2}", folded)]
        year = path_year or (years[-1] if years else 0)
        if not year:
            continue
        for name, month in MONTHS.items():
            match = re.search(rf"\b(\d{{1,2}})\s+(?:DE\s+)?{name}\b", folded)
            if match:
                try:
                    return datetime(year, month, int(match.group(1)))
                except ValueError:
                    pass
        numeric = re.search(
            r"\b(\d{1,2})[-/.](\d{1,2})(?:[-/.](\d{2,4}))?\b", numeric_text
        )
        if numeric:
            year_value = int(numeric.group(3)) if numeric.group(3) else year
            if year_value < 100:
                year_value += 2000
            if path_year:
                year_value = path_year
            try:
                return datetime(year_value, int(numeric.group(2)), int(numeric.group(1)))
            except ValueError:
                pass
    if path_month:
        day = _filename_day(source.stem)
        if day:
            try:
                return datetime(path_year, path_month, day)
            except ValueError:
                return None
    return None


def _cell_datetime(value: Any, default_date: datetime | None) -> datetime | None:
    if isinstance(value, datetime):
        if (
            default_date
            and value.year != default_date.year
            and abs((value.date() - default_date.date()).days) > 2
        ):
            return datetime.combine(default_date.date(), value.time())
        return value.replace(tzinfo=None)
    if isinstance(value, date):
        if (
            default_date
            and value.year != default_date.year
            and abs((value - default_date.date()).days) > 2
        ):
            return default_date.replace(hour=0, minute=0, second=0)
        return datetime.combine(value, time())
    if isinstance(value, time):
        return datetime.combine(default_date.date(), value) if default_date else None
    if isinstance(value, (int, float)):
        numeric = int(value)
        if 0 <= float(value) < 1 and default_date:
            seconds = round(float(value) * 24 * 60 * 60)
            return default_date + timedelta(seconds=seconds)
        return _clock_from_digits(str(numeric), default_date)
    text = _clean_text(value).upper()
    if not text or "TURNO" in text or text in {"...", "....", "*****", "####"}:
        return None
    cleaned = re.sub(r"\s+", " ", text.replace(".", ":"))
    match = re.search(
        r"\b(\d{1,2})[:/-](\d{1,2})[:/-](\d{2,4})\s+(\d{1,2})[:.](\d{2})\b",
        text,
    )
    if match:
        year = int(match.group(3))
        if year < 100:
            year += 2000
        try:
            parsed = datetime(
                year,
                int(match.group(2)),
                int(match.group(1)),
                int(match.group(4)),
                int(match.group(5)),
            )
            if (
                default_date
                and parsed.year != default_date.year
                and abs((parsed.date() - default_date.date()).days) > 2
            ):
                return datetime.combine(default_date.date(), parsed.time())
            return parsed
        except ValueError:
            return None
    clock = re.search(r"\b(\d{1,2})[:.](\d{2})\b", cleaned)
    if clock and default_date:
        try:
            return default_date.replace(hour=int(clock.group(1)), minute=int(clock.group(2)), second=0)
        except ValueError:
            return None
    return _clock_from_digits(text, default_date)


def _clock_from_digits(value: str, default_date: datetime | None) -> datetime | None:
    digits = re.sub(r"\D", "", value)
    if not default_date or len(digits) not in {3, 4}:
        return None
    hour = int(digits[:-2])
    minute = int(digits[-2:])
    if hour > 23 or minute > 59:
        return None
    return default_date.replace(hour=hour, minute=minute, second=0)


def _source_period(path: Path) -> tuple[int, int]:
    folded_parts = [fold_text(part) for part in path.parts]
    year = 0
    for part in reversed(folded_parts):
        years = re.findall(r"(?<!\d)(20\d{2})(?!\d)", part)
        if len(years) == 1:
            year = int(years[0])
            break
    month = 0
    for part in reversed(folded_parts[:-1]):
        match = re.match(r"(0?[1-9]|1[0-2])\D", part)
        if match:
            month = int(match.group(1))
            break
        month = next((number for name, number in MONTHS.items() if name in part), month)
        if month:
            break
    if not month:
        month = next((number for name, number in MONTHS.items() if name in folded_parts[-1]), 1)
    return year, month


def _empty_source_reason(source: Path) -> str:
    try:
        suffix = source.suffix.lower()
        snippets: list[str] = []
        if suffix == ".xlsx":
            workbook = openpyxl.load_workbook(source, read_only=True, data_only=True)
            try:
                for worksheet in workbook.worksheets:
                    for row_number, row in enumerate(worksheet.iter_rows(values_only=True)):
                        snippets.extend(_clean_text(value) for value in row if _clean_text(value))
                        if row_number >= 50:
                            break
            finally:
                workbook.close()
        elif suffix == ".xls":
            workbook = xlrd.open_workbook(source, on_demand=True)
            try:
                for worksheet in workbook.sheets():
                    for row in range(min(worksheet.nrows, 50)):
                        snippets.extend(
                            _clean_text(worksheet.cell_value(row, column))
                            for column in range(worksheet.ncols)
                            if _clean_text(worksheet.cell_value(row, column))
                        )
            finally:
                workbook.release_resources()
        elif suffix == ".docx":
            document = Document(source)
            snippets.extend(paragraph.text for paragraph in document.paragraphs[:50])
            for table in document.tables[:4]:
                for row in table.rows[:50]:
                    snippets.extend(cell.text for cell in row.cells)
        if "SIN NOVEDAD" in fold_text(" ".join(snippets)):
            return "sin_novedad"
    except Exception:
        return "archivo_danado"
    return "sin_registros_extraibles"


def _filename_day(value: str) -> int | None:
    match = re.search(r"(?:^|\D)([0-3]?\d)(?:\D|$)", value)
    if not match:
        return None
    day = int(match.group(1))
    return day if 1 <= day <= 31 else None


def _column(headers: list[str], *labels: str) -> int | None:
    return next(
        (index for index, value in enumerate(headers) if any(label in value for label in labels)),
        None,
    )


def _value(row: list[Any], index: int | None) -> Any:
    return row[index] if index is not None and index < len(row) else None


def _cell(row: list[Any], index: int | None) -> str:
    return _clean_text(_value(row, index))


def _clean_text(value: Any) -> str:
    return " ".join(str(value or "").split())


def _clean(value: str | None) -> str | None:
    value = (value or "").strip()
    return value if value and value not in {"-", "–", "...", "...."} else None


def _increment(values: dict[str, int], key: str, amount: int = 1) -> None:
    values[key] = values.get(key, 0) + amount


def _looks_sequence(value: Any) -> bool:
    text = _clean_text(value)
    return bool(re.fullmatch(r"\d{1,3}", text))


def _looks_clock(value: Any) -> bool:
    return isinstance(value, (datetime, time)) or bool(
        re.search(r"\b\d{1,2}[:.]\d{2}\b", _clean_text(value))
    )


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while chunk := handle.read(1024 * 1024):
            digest.update(chunk)
    return digest.hexdigest()
